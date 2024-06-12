# -*- coding: utf-8 -*-
"""
Started on Thu Feb 15 11:12:55 2024

@author: boris.sorokin@skao.int

This is a DB generator GUI with QT interface. Since it is supposed to be under GPL, should be fine
Version History:
v0.1a - Initial release
v0.1b - Fixed the bug in saving whole DB as a word function (incorrect evaluation of station frequency boundaries) and added VLBI Key. Also added an option to save the DB as SQLite DB.
v0.2a - Added a Wikidata query to the SQLite DB and Site Link Wizard to Link Wikidata stations with ITU ones
v0.2b - Fixed bugs with incorrect js build in Link Wizard

IDE used: VSCode with enviroment set and controlled by Anaconda
"""
import sys
import pyodbc
import sqlite3
import os
import base64
import csv
import docx
from docx.enum.section import WD_ORIENT
from SPARQLWrapper import SPARQLWrapper, JSON
import numpy as np

from PyQt5.QtWebEngineWidgets import QWebEngineView

from PyQt5.QtWidgets import (QApplication, QMainWindow, QDesktopWidget, QWidget, QPushButton, QFileDialog, QLabel,
                             QMessageBox, QGridLayout, QGroupBox, QDialog, QTableWidget, QTableWidgetItem, QCheckBox,
                             QHBoxLayout, QProgressDialog, QProgressBar, QListWidget, QSpacerItem, QSizePolicy, 
                             QListWidgetItem, QStackedLayout)

from PyQt5.QtGui import QIcon, QPixmap, QDesktopServices, QPainter, QColor

from PyQt5.QtCore import Qt, QParallelAnimationGroup, QPropertyAnimation, QRect, QEventLoop, QEasingCurve, QUrl, QTimer, QRectF

class MainApp(QMainWindow):
    """
    This is the main window with database selector, interactive database, and saving capability. 
    It serves as the entry point to all other classes below.
    """
    def __init__(self):
        """Initialize the main application window."""
        super().__init__()
        self.dbConnection = None
        self.interactive_database = None
        self.desired_width = 1280
        self.desired_height = 720
        self.initUI()

    def initUI(self):
        """Set up the user interface for the main application window."""
        # Set the window properties
        self.setWindowTitle('IAU CPS RAS Database Tool')
        self.setWindowIcon(QIcon('cps-logo-mono.ico'))

        # Center the window on the screen or parent
        self.centerWindow(self)
        self.animateOpening(self, self.desired_width, self.desired_height)

        # Create menubar
        menuBar = self.menuBar()
        helpMenu = menuBar.addMenu('Help')
        
        # Create about button in the menubar
        aboutAction = helpMenu.addAction('About')
        aboutAction.triggered.connect(self.show_about)

        # Create a group box for ITU Database tools
        self.ituToolsGroup = QGroupBox('ITU Database Tools')
        gridLayoutITU = QGridLayout()

        # Create a button for selecting the database file
        self.button_open = QPushButton('Select ITU database to import', self)
        self.button_open.clicked.connect(self.database_select)
        # Place the button in the grid layout
        gridLayoutITU.addWidget(self.button_open, 0, 0)

        # Create a status indicator for selecting the file
        self.statusLight_open = QLabel(self)
        self.statusLight_open.setFixedSize(20, 20)
        self.updateStatusLight(self.statusLight_open,
                               False, 'Database not selected.')
        gridLayoutITU.addWidget(self.statusLight_open, 0, 1)
        gridLayoutITU.setColumnStretch(1, 0)

        # Create a button for connecting to the selected database
        self.button_connect = QPushButton('Connect to selected database', self)
        self.button_connect.clicked.connect(self.database_connect)
        # Place the button in the grid layout
        gridLayoutITU.addWidget(self.button_connect, 1, 0)
        self.button_connect.setEnabled(False)
        self.button_connect.setToolTip('Select a database first.')

        # Create a status indicator for selecting the file
        self.statusLight_connect = QLabel(self)
        self.statusLight_connect.setFixedSize(20, 20)
        self.updateStatusLight(self.statusLight_connect,
                               False, 'Database not connected.')
        gridLayoutITU.addWidget(self.statusLight_connect, 1, 1)
        gridLayoutITU.setColumnStretch(1, 0)

        # Create a button for showing the interactive list
        self.button_show_list = QPushButton(
            'Show the list of radio astronomy stations', self)
        self.button_show_list.clicked.connect(self.interactive_database_show)
        # Place the button in the grid layout
        gridLayoutITU.addWidget(self.button_show_list, 2, 0)
        self.button_show_list.setEnabled(False)
        self.button_show_list.setToolTip('Connect a database first.')

        self.ituToolsGroup.setLayout(gridLayoutITU)

        # Create a group box for ITU Database tools
        self.exportToolsGroup = QGroupBox('Export Tools')
        gridLayoutExport = QGridLayout()

        self.button_export_csv = QPushButton(
            'Export all data as CSV', self)
        self.button_export_csv.clicked.connect(self.save_csv)
        self.button_export_csv.setEnabled(False)
        self.button_export_csv.setToolTip('Connect ITU database first.')
        gridLayoutExport.addWidget(self.button_export_csv, 0, 0)

        self.button_export_word = QPushButton(
            'Export all data as DOCX', self)
        self.button_export_word.clicked.connect(self.save_word)
        self.button_export_word.setEnabled(False)
        self.button_export_word.setToolTip('Connect ITU database first.')
        gridLayoutExport.addWidget(self.button_export_word, 0, 1)
        
        self.button_export_SQLite = QPushButton(
            'Export all data as SQLite DB', self)
        self.button_export_SQLite.clicked.connect(self.save_DB)
        self.button_export_SQLite.setEnabled(False)
        self.button_export_SQLite.setToolTip('Connect ITU database first.')
        gridLayoutExport.addWidget(self.button_export_SQLite, 0, 2)

        self.button_runSiteLinkWizard = QPushButton(
            'Run Site Link Wizard', self)
        self.button_runSiteLinkWizard.clicked.connect(self.run_site_link_wizard)
        gridLayoutExport.addWidget(self.button_runSiteLinkWizard, 1, 0)

        self.exportToolsGroup.setLayout(gridLayoutExport)

        # Setting central widget
        centralWidget = QWidget()
        self.setCentralWidget(centralWidget)

        # Main layout for the widget
        mainLayout = QGridLayout(centralWidget)
        mainLayout.addWidget(self.ituToolsGroup, 0, 0)
        mainLayout.addWidget(self.exportToolsGroup, 1, 0)

        self.statusBar().showMessage('    Ready')
        self.statusBar().setStyleSheet("""
            QStatusBar {
                background-color: #404040;
                color: white;
                border: 2px solid black;
            }
        """)
        self.show()
        self.setFocus()
        self.raise_()
        self.activateWindow()

    def database_select(self):
        # Selecting database callback
        options = QFileDialog.Options()
        self.database_file_name, _ = QFileDialog.getOpenFileName(
            self, "Select ITU Database File", "", "MDB Files (*.mdb)", options=options)
        if self.database_file_name:
            self.statusBar().showMessage(
                f'    Database file selected: {self.database_file_name}')
            self.updateStatusLight(self.statusLight_open,
                                   True, 'Database selected')
            self.button_connect.setEnabled(True)
            self.button_connect.setToolTip(None)
        else:
            self.updateStatusLight(self.statusLight_open,
                                   False, 'Database not selected')
            self.button_connect.setEnabled(False)
            self.button_connect.setToolTip('Select a database first.')

            self.button_export_csv.setEnabled(False)
            self.button_export_csv.setToolTip('Select a database first.')

            self.button_export_word.setEnabled(False)
            self.button_export_word.setToolTip('Select a database first.')            
            
            self.button_export_SQLite.setEnabled(False)
            self.button_export_SQLite.setToolTip('Select a database first.')

    def database_connect(self):
        # Attempt to connect to the selected database
        try:
            connectionString = f'DRIVER={{Microsoft Access Driver (*.mdb, *.accdb)}};DBQ={self.database_file_name}'
            try:
                self.dbConnection.close() # type: ignore
            except:
                pass

            self.dbConnection = pyodbc.connect(connectionString)
            self.updateStatusLight(
                self.statusLight_connect, True, 'Database connected')
            self.statusBar().showMessage('    Database connected. Checking version...')
            SQL = "SELECT d_create, comment FROM srs_ooak"
            rows = self.parse_database(SQL)[0]
            self.database_date = rows[0]
            self.database_version = rows[1][0:7]
            self.statusBar().showMessage(
                f'Connected to database {self.database_version} published on {self.database_date.date()}')
            self.button_show_list.setEnabled(True)
            self.button_show_list.setToolTip(None)

            self.button_export_csv.setEnabled(True)
            self.button_export_csv.setToolTip(None)

            self.button_export_word.setEnabled(True)
            self.button_export_word.setToolTip(None)
            
            self.button_export_SQLite.setEnabled(True)
            self.button_export_SQLite.setToolTip(None)
        except Exception as e:
            QMessageBox.critical(self, "Database Connection Error",
                                 f"An error occurred while connecting to the database:\n{e}")
            self.updateStatusLight(
                self.statusLight_connect, True, 'Database connection error')
            self.statusBar().showMessage('    Database connection error')
            self.button_show_list.setToolTip('Connect a database first.')

    def interactive_database_show(self):
        self.animateClosing(self)
        self.showMinimized()
        self.setEnabled(False)
        self.interactive_database = InteractiveDatabase(self)

    def parse_database(self, SQL):
        """Parse the database and display results in a new window and save to a Word document."""
        rows = []
        if self.dbConnection:
            try:
                cursor = self.dbConnection.cursor()
                cursor.execute(SQL)
                rows = cursor.fetchall()
                cursor.close()
            except Exception as e:
                QMessageBox.critical(self,
                                     "Database Error", f"Error parsing database: {e}")
        return rows

    def show_about(self):
        # Show about window.
        aboutDialog = AboutDialog(self)
        aboutDialog.exec_()

    def save_csv(self):
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save as CSV", f"RAS_DB_FULL_CSV_{self.database_version}_{self.database_date.date()}", "CSV Files (*.csv)")
        if filePath:
            if os.path.basename(filePath) == 'geographical-areas.csv':
                QMessageBox.critical(
                    self, "Error", "This file is an important app file and cannot be overwritten.")
                return

            SQL = 'SELECT ntc_id, adm, ctry, stn_name, long_dec, lat_dec FROM com_el WHERE ntc_type=\'R\' ORDER BY adm asc, stn_name asc;'
            com_el_rows = self.parse_database(SQL)
            station_number = len(com_el_rows)

            fields = ['Notice ID', 'Administration', 'Region/Location', 'Station name',
                      'Longitude', 'Latitude', 'Longitude Degrees',
                      'Longitude East/West', 'Longitude minutes', 'Longitude seconds',
                      'Latitude Degrees', 'Latitude North/South', 'Latitude minutes',
                      'Latitude seconds', 'Elevation minimum', 'Elevation maximum',
                      'Azimuth from', 'Azimuth to', 'Beam name', 'Antenna pattern ID',
                      'Antenna pattern Name', 'Centre frequency, MHz',
                      'Group ID', 'Noise Temp, K', 'Frequency minimum, MHz',
                      'Frequency maximum, MHz', 'Date brought into use', 'Date received',
                      'IFIC no (wic_no)', 'Date updated', 'VLBI Support code']

            with open(filePath, 'w', newline='') as file:
                csv_writer = csv.writer(file, delimiter=',')
                csv_writer.writerow(fields)
                for index in range(0, station_number):
                    ntc_id = com_el_rows[index][0]
                    SQL = 'SELECT long_deg, long_ew, long_min, long_sec, lat_deg, lat_ns, lat_min, lat_sec, elev_min, elev_max, azm_fr, azm_to FROM e_stn WHERE ntc_id=' + \
                        str(ntc_id)+';'
                    e_stn_rows = self.parse_database(SQL)
                    SQL = 'SELECT beam_name, pattern_id FROM e_ant WHERE ntc_id=' + \
                        str(ntc_id)+';'
                    e_ant_rows = self.parse_database(SQL)

                    beam_number = len(e_ant_rows)
                    beam_names = []
                    ant_id = []
                    ant_names = []

                    for subindex_beam in range(0, beam_number):
                        beam_names.append(e_ant_rows[subindex_beam][0])
                        ant_id.append(e_ant_rows[subindex_beam][1])
                        if ant_id[-1] == None:
                            SQL = 'SELECT attch_e, ant_diam FROM e_ant WHERE ntc_id=' + \
                                str(ntc_id)+';'
                            local_ant_rows = self.parse_database(SQL)
                            if local_ant_rows[subindex_beam][1] == None:
                                ant_names.append('NonTypical, see attachment {} to the relevant IFIC for details.'.format(
                                    local_ant_rows[subindex_beam][0]))
                            else:
                                ant_names.append('NonTypical, submitted diameter is {} meters, see attachment {} to the relevant IFIC for details.'.format(
                                    local_ant_rows[subindex_beam][1], local_ant_rows[subindex_beam][0]))
                        else:
                            SQL = 'SELECT pattern FROM ant_type WHERE pattern_id=' + \
                                str(ant_id[-1])+';'
                            ant_type_rows = self.parse_database(SQL)
                            try:
                                ant_names.append(ant_type_rows[0][0])
                            except:
                                SQL = 'SELECT attch_e, ant_diam FROM e_ant WHERE ntc_id=' + \
                                    str(ntc_id)+';'
                                local_ant_rows = self.parse_database(SQL)
                                if local_ant_rows[subindex_beam][1] == None:
                                    ant_names.append('NonTypical, see attachment {} to the relevant IFIC for details.'.format(
                                        local_ant_rows[subindex_beam][0]))
                                else:
                                    ant_names.append('NonTypical, submitted diameter is {} meters, see attachment {} to the relevant IFIC for details.'.format(
                                        local_ant_rows[subindex_beam][1], local_ant_rows[subindex_beam][0]))

                        SQL = 'SELECT grp_id, noise_t, freq_min, freq_max, d_inuse, d_rcv, wic_no, d_upd, ra_stn_type FROM grp WHERE (ntc_id='+str(
                            ntc_id)+'and beam_name=\''+beam_names[subindex_beam]+'\''+');'
                        grp_rows = self.parse_database(SQL)

                        SQL = 'SELECT freq_mhz FROM freq WHERE (ntc_id='+str(
                            ntc_id)+'and beam_name=\''+beam_names[subindex_beam]+'\''+');'
                        freq_rows = self.parse_database(SQL)

                        group_number = len(grp_rows)

                        for subindex_group in range(0, group_number):
                            if subindex_beam == 0 and subindex_group == 0:
                                next_line = list(com_el_rows[index])+list(e_stn_rows[0])+list([beam_names[subindex_beam]])+list([ant_id[subindex_beam]])+list(
                                    [ant_names[subindex_beam]])+list([str(freq_rows[subindex_group][0])])+list(grp_rows[subindex_group])
                                csv_writer.writerow(next_line)
                            else:
                                next_line = list([''])*18+list([beam_names[subindex_beam]])+list([ant_id[subindex_beam]])+list(
                                    [ant_names[subindex_beam]])+list([str(freq_rows[subindex_group][0])])+list(grp_rows[subindex_group])
                                csv_writer.writerow(next_line)

    def save_word(self):
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save as DOCX", f"RAS_DB_FULL_DOCX_{self.database_version}_{self.database_date.date()}", "Word Files (*.docx)")
        if filePath:
            doc = docx.Document()
            section = doc.sections[0]
            section = doc.sections[0]
            section.orientation = WD_ORIENT.PORTRAIT
            doc.add_heading(
                "Annex 1. The list of radio astronomy stations known to the IAU CPS", level=1)
            doc.add_paragraph("This list is based on ITU-R IFIC database.")

            SQL = 'SELECT ntc_id, adm, ctry, stn_name, long_dec, lat_dec FROM com_el WHERE ntc_type=\'R\' ORDER BY adm asc, stn_name asc;'
            com_el_rows = self.parse_database(SQL)
            station_number = len(com_el_rows)

            progressDialog = QProgressDialog(
                "Operation in progress...", "Cancel", 0, station_number, self)
            progressDialog.setWindowTitle("Saving...")
            progressDialog.setWindowModality(Qt.WindowModal)
            progressDialog.setCancelButton(None)
            progressDialog.show()
                   
            country_codes_to_names=self.load_country_codes()

            for index, row in enumerate(com_el_rows):
                progressDialog.setLabelText(
                    f"Now populating station {index+1} of {station_number}")
                ntc_id = row[0]
                try:
                    admin_name = country_codes_to_names[row[1]]
                except:
                    admin_name = 'Unknown'
                try:
                    country_name = country_codes_to_names[row[2]]
                except:
                    country_name = 'Unknown'
                station_name = row[3]
                station_longitude = com_el_rows[index][4]
                station_latitude = com_el_rows[index][5]

                SQL = f'SELECT elev_min, ant_alt FROM e_stn WHERE ntc_id={ntc_id};'
                e_stn_rows = self.parse_database(SQL)

                station_min_elevation = e_stn_rows[0][0]
                if station_min_elevation == None:
                    station_min_elevation = 'N/A'
                    
                station_antenna_altitude = e_stn_rows[0][1]
                if station_antenna_altitude == None:
                    station_antenna_altitude = 'N/A'

                SQL = 'SELECT beam_name, ant_diam, gain FROM e_ant WHERE ntc_id=' + \
                    str(ntc_id)+';'
                e_ant_rows = self.parse_database(SQL)

                beam_number = len(e_ant_rows)

                beam_names = []
                noise_temp = []
                ant_diameter = []
                ant_gain = []
                freq_min = []
                freq_max = []
                vlbi_type = []                

                for subindex_beam in range(0, beam_number):
                    beam_names.append(e_ant_rows[subindex_beam][0])
                    if e_ant_rows[subindex_beam][1] == None:
                        ant_diameter.append('N/A')
                    else:
                        ant_diameter.append(e_ant_rows[subindex_beam][1])

                    if e_ant_rows[subindex_beam][2] == None:
                        ant_gain.append('N/A')
                    else:
                        ant_gain.append(e_ant_rows[subindex_beam][2])

                    SQL = f"SELECT noise_t, freq_min, freq_max, ra_stn_type FROM grp WHERE (ntc_id={ntc_id} and beam_name='{beam_names[subindex_beam]}');"
                    grp_rows = self.parse_database(SQL)
                    
                    for group_index, group in enumerate(grp_rows):
                        noise_temp.append(grp_rows[group_index][0])
                        freq_min.append(grp_rows[group_index][1])
                        freq_max.append(grp_rows[group_index][2])
                        vlbi_type.append(grp_rows[group_index][3])

                station_freq_min = min(freq_min)
                station_freq_max = max(freq_max)

                # Flushing all gathered data to the document
                doc.add_heading(f'Station "{station_name}"', level=2)
                doc.add_heading('Overview', level=3)
                doc.add_paragraph(f'Station number: {index+1}')
                doc.add_paragraph(
                    f'Responsible administration: "{admin_name}"')
                doc.add_paragraph(f'Country/region location: "{country_name}"')
                doc.add_paragraph(f'Station short name: "{station_name}')
                doc.add_paragraph('Station long: name "N/A"')
                doc.add_paragraph('Station type: "N/A"')
                doc.add_paragraph(
                    f'Station longitude [deg]: "{station_longitude}"')
                doc.add_paragraph(
                    f'Station latitude [deg]: "{station_latitude}"')
                doc.add_paragraph(f'Station altitude (AMSL) "{station_antenna_altitude}"')
                doc.add_paragraph(
                    f'Minimum elevation [deg]: "{station_min_elevation}"')
                doc.add_paragraph('Operational "N/A"')
                doc.add_paragraph('Used for science "N/A"')
                doc.add_paragraph(
                    f'Minimum Station Frequency [MHz]: "{station_freq_min} MHz"')
                doc.add_paragraph(
                    f'Maximum Station Frequency [MHz]: "{station_freq_max} MHz"')
                doc.add_paragraph('Contact (website) "N/A"')
                doc.add_paragraph('Contact (address) "N/A"')
                doc.add_paragraph('Contact (phone) "N/A"')
                doc.add_paragraph('Contact (e-mail) "N/A"')

                doc.add_heading('Antenna information', level=3)

                for beam_index in range(0, beam_number):
                    doc.add_heading(f'Antenna #{beam_index+1}', level=4)
                    doc.add_paragraph('Feed/Rx height above ground [m] "N/A"')
                    doc.add_paragraph(
                        f'Noise temparature [K]: "{noise_temp[beam_index]}"')
                    doc.add_paragraph(
                        f'Antenna diameter [m]: "{ant_diameter[beam_index]}"')
                    doc.add_paragraph(
                        f'Maximum antenna gain [dBi]: "{ant_gain[beam_index]}"')
                    doc.add_paragraph(
                        f'Minimim antenna frequency [MHz]: "{freq_min[beam_index]}"')
                    doc.add_paragraph(
                        f'Maximum antenna frequency [MHz]: "{freq_max[beam_index]}"')
                    doc.add_paragraph('Cryocooled: "N/A"')
                    doc.add_paragraph('Supports RAS mode continuum: "N/A"')
                    doc.add_paragraph('Supports RAS mode spectroscopy: "N/A"')
                    
                    # 'S' stands for single dish and 'V' for VLBI
                    if(vlbi_type[beam_index]=='V'):
                        doc.add_paragraph('Supports RAS mode VLBI: "Yes"')
                    elif(vlbi_type[beam_index]=='S'):
                        doc.add_paragraph('Supports RAS mode VLBI: "No"')
                    else:                        
                        doc.add_paragraph('Supports RAS mode VLBI: "N/A"')

                doc.add_page_break()
                progressDialog.setValue(index+1)
                

            try:
                doc.save(filePath)
            except Exception as e:
                QMessageBox.critical(self, "Docx saving Error",
                                     f"An error occurred while preparing the docx file:\n{e}")
            progressDialog.close()
    def save_DB(self):
        def create_database(filePath):
            if os.path.exists(filePath):
                os.remove(filePath)
            conn = sqlite3.connect(filePath, timeout=90)
            cursor = conn.cursor()
            # SQL commands to create tables
            create_table_stations = """
            CREATE TABLE "Stations" (
                "CPS Station ID"	INTEGER NOT NULL UNIQUE,
                "Country"	TEXT,
                "Short Name"	TEXT NOT NULL,
                "Long Name"	TEXT,
                "Type"	TEXT CHECK("Type" IN ('single dish', 'array', 'mixed', 'unknown')),
                "Station longitude [deg]"	NUMERIC,
                "Station latitude [deg]"	NUMERIC,
                "Station altitude (amsl) [m]"	NUMERIC,
                "Operational"	INTEGER,
                "Used for science"	INTEGER,
                "Min station frequency [MHz]"	NUMERIC,
                "Max station frequency [MHz]"	NUMERIC,
                "Contact / Website"	TEXT,
                "Contact / Address"	TEXT,
                "Contact / Phone"	TEXT,
                "Contact / Email"	TEXT,
                "Registered at ITU"	INTEGER,
                "ITU Notice ID"	INTEGER,
                "ITU responsible Administration"	TEXT,
                PRIMARY KEY("CPS Station ID" AUTOINCREMENT)
            );
            """
            cursor.execute(create_table_stations)
            create_table_antennas = """
            CREATE TABLE "Antennas" (
            	"CPS Station ID"	INTEGER,
            	"CPS Antenna ID"	INTEGER NOT NULL UNIQUE,
            	"Antenna longitude [deg]"	NUMERIC,
            	"Antenna latitude [deg]"	NUMERIC,
            	"Antenna altitude (WGS84) [m]"	NUMERIC,
            	"Antenna altitude (amsl) [m]"	NUMERIC,
            	"Feed/Rx height above ground [m]"	NUMERIC,
            	"Antenna diameter [m]"	NUMERIC,
            	"Minimum elevation [deg]"	NUMERIC,
            	"Minimum frequency [MHz]"	NUMERIC,
            	"Maximum frequency [MHz]"	NUMERIC,
            	FOREIGN KEY("CPS Station ID") REFERENCES "Stations"("CPS Station ID"),
            	PRIMARY KEY("CPS Antenna ID" AUTOINCREMENT)
            );
            """
            cursor.execute(create_table_antennas)
            create_table_frequency_bands = """
            CREATE TABLE "Frequency_Bands" (
            	"CPS Station ID"	INTEGER,
            	"CPS Antenna ID"	INTEGER,
            	"CPS Band ID"	INTEGER NOT NULL UNIQUE,
            	"Band start [MHz]"	NUMERIC,
            	"Band stop [MHz]"	NUMERIC,
            	"Antenna eff. Area [m^2]"	NUMERIC,
            	"Cryo-cooled"	INTEGER,
            	"Polarisation"	TEXT,
            	"Supports RAS mode continuum"	INTEGER,
            	"Supports RAS mode spectroscopy"	INTEGER,
            	"Supports RAS mode VLBI"	INTEGER,
            	"Noise temperature [K]"	NUMERIC,
            	FOREIGN KEY("CPS Station ID") REFERENCES "Stations"("CPS Station ID"),
            	FOREIGN KEY("CPS Antenna ID") REFERENCES "Antennas",
            	PRIMARY KEY("CPS Band ID" AUTOINCREMENT)
            );
            """
            cursor.execute(create_table_frequency_bands)
            create_table_wikidata="""
            CREATE TABLE IF NOT EXISTS wikidata (
                "CPS Wiki ID" INTEGER,
                Name TEXT,
                Country TEXT,
                "Station longitude [deg]"	NUMERIC,
                "Station latitude [deg]"	NUMERIC,  
                source TEXT,              
                "Linked ITU" INTEGER,
            	PRIMARY KEY("CPS Wiki ID")
            );
            """            
            cursor.execute(create_table_wikidata)
            return conn, cursor
        
        def process_stations(cursor_ITU, cursor_CPS, country_codes_to_names):
            """
            Processes each station from the ITU database and inserts it into the CPS database.
            """
            cursor_ITU.execute('SELECT ntc_id, adm, ctry, stn_name, long_dec, lat_dec FROM com_el WHERE ntc_type=\'R\' ORDER BY adm asc, stn_name asc;')
            com_el_rows = cursor_ITU.fetchall()
            station_number=len(com_el_rows)
            progressDialog = QProgressDialog(
                "Operation in progress...", "Cancel", 0, station_number, self)
            progressDialog.setWindowTitle("Saving...")
            progressDialog.setWindowModality(Qt.WindowModal)
            progressDialog.setCancelButton(None)
            progressDialog.show()           
            

            for index, row in enumerate(com_el_rows):
                country_name = country_codes_to_names.get(row[2], 'Unknown')
                progressDialog.setLabelText(
                    f"Now populating station {index+1} of {station_number}")
                cursor_CPS.execute('''
                    INSERT INTO stations ("ITU Notice ID", "ITU responsible Administration", "Country", "Short name",
                                          "Station longitude [deg]", "Station latitude [deg]", "Registered at ITU")
                    VALUES (?, ?, ?, ?, ?, ?, 1);
                ''', (row[0], row[1], country_name, row[3], row[4], row[5]))
                cps_station_id = cursor_CPS.lastrowid
                process_antennas(cursor_ITU, cursor_CPS, row[0], cps_station_id, row[4], row[5])
                progressDialog.setValue(index+1)
                
            progressDialog.close()
            
        def process_antennas(cursor_ITU, cursor_CPS, ntc_id, cps_station_id, long_dec, lat_dec):
            """
            Retrieves antenna data related to the station and processes frequency bands for each antenna.
            """
            cursor_ITU.execute(f"SELECT beam_name, ant_diam FROM e_ant WHERE ntc_id={ntc_id};")
            e_ant_rows = cursor_ITU.fetchall()

            cursor_ITU.execute(f"SELECT elev_min FROM e_stn WHERE ntc_id={ntc_id};")
            elev_min_result = cursor_ITU.fetchone()
            elev_min = elev_min_result[0] if elev_min_result else None

            frequency_ranges = []

            for beam_name, ant_diam in e_ant_rows:
                cursor_CPS.execute('''
                    INSERT INTO antennas ("CPS Station ID", "Antenna diameter [m]", "Minimum elevation [deg]", "Antenna longitude [deg]", "Antenna latitude [deg]")
                    VALUES (?, ?, ?, ?, ?);
                ''', (cps_station_id, ant_diam, elev_min, long_dec, lat_dec))
                cps_antenna_id = cursor_CPS.lastrowid
                freqs = process_frequency_bands(cursor_ITU, cursor_CPS, ntc_id, beam_name, cps_antenna_id, cps_station_id)
                frequency_ranges.extend(freqs)        
                min_freq = min(freqs)
                max_freq = max(freqs)
                cursor_CPS.execute('''
                    UPDATE antennas SET "Minimum frequency [MHz]" = ?, "Maximum frequency [MHz]" = ?
                    WHERE "CPS Antenna ID" = ?;
                ''', (min_freq, max_freq, cps_antenna_id))

            
            min_freq = min(frequency_ranges)
            max_freq = max(frequency_ranges)
            cursor_CPS.execute('UPDATE stations SET "Min station frequency [MHz]" = ?, "Max station frequency [MHz]" = ? WHERE "CPS Station ID" = ?;', (min_freq, max_freq, cps_station_id))
                

        def process_frequency_bands(cursor_ITU, cursor_CPS, ntc_id, beam_name, cps_antenna_id, cps_station_id):
            """
            Processes frequency bands for each antenna and inserts them into the CPS database.
            """
            cursor_ITU.execute(f"SELECT freq_min, freq_max, ra_stn_type, noise_t FROM grp WHERE ntc_id={ntc_id} AND beam_name='{beam_name}';")
            grp_rows = cursor_ITU.fetchall()

            frequency_ranges = []

            for freq_min, freq_max, ra_stn_type, noise_t in grp_rows:
                vlbi_key = 1 if ra_stn_type == 'V' else 0
                cursor_CPS.execute('''
                    INSERT INTO Frequency_Bands ("CPS Station ID", "CPS Antenna ID", "Band start [MHz]", "Band stop [MHz]", "Supports RAS mode VLBI", "Noise temperature [K]")
                    VALUES (?, ?, ?, ?, ?, ?);
                ''', (cps_station_id, cps_antenna_id, freq_min, freq_max, vlbi_key, noise_t))
                frequency_ranges.append(freq_min)
                frequency_ranges.append(freq_max)

            return frequency_ranges

        def add_wiki_data(cursor_CPS):
            sparql_query = """
            SELECT ?item ?itemLabel ?countryLabel ?coordinate_location WHERE {
            VALUES ?val {
                wd:Q184356
                wd:Q349772
            }
            ?item wdt:P31 ?val.
            SERVICE wikibase:label { bd:serviceParam wikibase:language "[AUTO_LANGUAGE],en,es,ja,ru". }
            OPTIONAL { ?item wdt:P625 ?coordinate_location. }
            OPTIONAL { ?item wdt:P17 ?country. }
            }
            ORDER BY (?itemLabel)
            LIMIT 1000
            """

            # Set up the SPARQL endpoint
            sparql = SPARQLWrapper("https://query.wikidata.org/sparql")
            sparql.setQuery(sparql_query)
            sparql.setReturnFormat(JSON)

            results = sparql.query().convert()
            data = results["results"]["bindings"] # type: ignore
            insert_query = """
            INSERT INTO wikidata (Name, Country, "Station longitude [deg]", "Station latitude [deg]", source)
            VALUES (?, ?, ?, ?, ?)
            """
            for item in data:
                name = item.get("itemLabel", {}).get("value", "") # type: ignore
                country = item.get("countryLabel", {}).get("value", "") # type: ignore
                coordinates = item.get("coordinate_location", {}).get("value", "") # type: ignore
                if coordinates:
                    longitude, latitude = coordinates.strip('Point()').split()
                else:
                    longitude, latitude = None, None
                source = item.get("item", {}).get("value", "") # type: ignore
                # push processed data to populate the wikidata DB
                cursor_CPS.execute(insert_query, (name, country, longitude, latitude, source))
                

                
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save as SQLite", f"CPS_RAS_DB_FULL_SQLite_{self.database_version}_{self.database_date.date()}", 
            "SQLite Database Files (*.db)")
        if filePath:
            if not filePath.endswith('.db'):
                filePath += '.db'
            try:
                conn, cursor = create_database(filePath)
                country_codes_to_names=self.load_country_codes()
                
                process_stations(self.dbConnection.cursor(), cursor, country_codes_to_names) # type: ignore
                add_wiki_data(cursor)
                
                conn.commit()
                conn.close()

                # Asking the user if they want to run the Site Link Wizard
                reply = QMessageBox.question(self, 'Run Site Link Wizard',
                                            'Would you like to run Site Link Wizard?',
                                            QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
                if reply == QMessageBox.Yes:                    
                    self.animateClosing(self)
                    self.showMinimized()
                    self.setEnabled(False)
                    self.interactive_database = SiteLinkWizard(filePath, self)
            except Exception as e:
                QMessageBox.critical(self, "DB saving Error",
                                     f"An error occurred while preparing the db file:\n{e}")  
            
    def run_site_link_wizard(self):
        options = QFileDialog.Options()
        filePath, _ = QFileDialog.getOpenFileName(self, "Select CPS Database File", "", "SQLite Files (*.db);;All Files (*)", options=options)
        if filePath:
            self.animateClosing(self)
            self.showMinimized()
            self.setEnabled(False)
            self.site_link_wizard = SiteLinkWizard(filePath, parent=self)
            self.site_link_wizard.show()
        
    def load_country_codes(self, filepath=None):
        """
        Loads country codes and their corresponding names from a CSV file to
        establish a link between ITU country codes and country names
        https://www.itu.int/en/ITU-R/terrestrial/fmd/Pages/geo_area_list.aspx 
        """
        country_codes_to_names = {}
        if filepath is None:
            filepath = 'geographical-areas.csv'
        with open(filepath, newline='') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                country_codes_to_names[row[0]] = row[1]
        return country_codes_to_names
    
    def updateStatusLight(self, widget, status, status_text):
        # Update the status light color based on whether connection was successfully established
        if status:
            widget.setStyleSheet(
                "QLabel { background-color: green; border-radius: 10px; }")
        else:
            widget.setStyleSheet(
                "QLabel { background-color: red; border-radius: 10px; }")

        self.connectionStatus = status_text
        widget.setToolTip(status_text)

    def centerWindow(self, window, parent=None):
        # Center the window on the screen or parent window
        if parent:
            parentGeometry = parent.frameGeometry()
            centerPoint = parentGeometry.center()
            window.move(centerPoint - window.rect().center())
        else:
            screenGeometry = QDesktopWidget().screenGeometry()
            centerPoint = screenGeometry.center()
            window.move(centerPoint - window.rect().center())

    def animateOpening(self, window, desired_width, desired_height, parent=None):
        # This function animates opening of the window

        # Block interaction
        window.setEnabled(False)

        if parent:
            centerPoint = parent.frameGeometry().center()
        else:
            desktopWidget = QDesktopWidget()
            centerPoint = desktopWidget.availableGeometry(
                desktopWidget.primaryScreen()).center()

        startRect = QRect(0, 0, int(desired_width * 0.1),
                          int(desired_height * 0.1))
        startRect.moveCenter(centerPoint)
        endRect = QRect(0, 0, desired_width, desired_height)
        endRect.moveCenter(centerPoint)

        # Create and configure the geometry animation
        window.animation_geometry_opening = QPropertyAnimation(
            window, b"geometry")
        window.animation_geometry_opening.setDuration(
            600)  # Duration of animation in ms
        window.animation_geometry_opening.setStartValue(startRect)
        window.animation_geometry_opening.setEndValue(endRect)
        window.animation_geometry_opening.setEasingCurve(
            QEasingCurve.InOutCubic)

        # Create and configure the opacity animation
        window.animation_opacity_opening = QPropertyAnimation(
            window, b"windowOpacity")
        window.animation_opacity_opening.setDuration(500)
        window.animation_opacity_opening.setStartValue(-0.5)
        window.animation_opacity_opening.setEndValue(1.0)
        window.animation_opacity_opening.setEasingCurve(
            QEasingCurve.InOutCubic)

        # Grouping animations
        window.animation_group_opening = QParallelAnimationGroup()
        window.animation_group_opening.addAnimation(
            window.animation_geometry_opening)
        window.animation_group_opening.addAnimation(
            window.animation_opacity_opening)

        window.animation_geometry_opening.finished.connect(
            lambda: window.setEnabled(True))
        window.animation_group_opening.start()

    def animateClosing(self, window):
        # This function animates closing of the window.

        # Block interactions
        window.setEnabled(False)

        # Get current size and position
        currentRect = window.geometry()
        centerPoint = currentRect.center()

        # Calculate the final rectangle (10% of the current size)
        endRect = QRect(0, 0, int(currentRect.width() * 0.1),
                        int(currentRect.height() * 0.1))
        endRect.moveCenter(centerPoint)

        # Create and configure the geometry animation
        window.animation_geometry_closing = QPropertyAnimation(
            window, b"geometry")
        window.animation_geometry_closing.setDuration(
            500)  # Duration in milliseconds
        window.animation_geometry_closing.setStartValue(currentRect)
        window.animation_geometry_closing.setEndValue(endRect)
        window.animation_geometry_closing.setEasingCurve(
            QEasingCurve.InOutCubic)

        # Create and configure the opacity animation
        window.animation_opacity_closing = QPropertyAnimation(
            window, b"windowOpacity")
        window.animation_opacity_closing.setDuration(500)
        window.animation_opacity_closing.setStartValue(1.0)
        window.animation_opacity_closing.setEndValue(-0.5)
        window.animation_opacity_closing.setEasingCurve(
            QEasingCurve.InOutCubic)

        # Grouping animations
        window.animation_group_closing = QParallelAnimationGroup()
        window.animation_group_closing.addAnimation(
            window.animation_geometry_closing)
        window.animation_group_closing.addAnimation(
            window.animation_opacity_closing)

        loop = QEventLoop()
        window.animation_geometry_closing.finished.connect(window.hide)
        window.animation_geometry_closing.finished.connect(loop.quit)
        window.animation_group_closing.start()
        loop.exec_()

    def closeEvent(self, event):
        # Check if the database connection is active and close it before exiting
        self.animateClosing(self)

        if self.dbConnection:
            try:
                self.dbConnection.close()
                self.statusBar().showMessage('    Database connection closed. App will soon close')
            except Exception:
                self.statusBar().showMessage(
                    '    Database connection was already closed. App will soon close')
        event.accept()


class AboutDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent

        self.initUI()

    def initUI(self):
        self.setWindowTitle("About")
        self.setWindowIcon(QIcon('cps-logo-mono.ico'))
        self.setWindowFlags(self.windowFlags(
        ) & ~Qt.WindowContextHelpButtonHint | Qt.WindowCloseButtonHint)
        self.parent.animateOpening(self, desired_width=640, desired_height=360) # type: ignore
        self.setModal(True)

        layout = QGridLayout()

        imageLabel = QLabel(self)
        pixmap = QPixmap("CPS_Logo_Col.png")
        pixmap = pixmap.scaled(
            300, 300, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        imageLabel.setPixmap(pixmap)
        layout.addWidget(imageLabel, 0, 0, 2, 1)

        textLabel1 = QLabel("This tool helps with importing ITU database for IAU CPS RAS database.\n\n"
                            "Program version: v0.2a\n\n"
                            "This version introduced Site Link Wizard to connect wikidata entries to ITU\n\n", self)
        textLabel1.setWordWrap(True)
        layout.addWidget(textLabel1, 0, 1)

        textLabel2 = QLabel("For more information, contact via either:<br><br>"
                            "<a href='mailto:ras.database@cps.iau.org'>ras.database@cps.iau.org</a><br><br>"
                            "<a href='mailto:boris.sorokin@skao.int'>boris.sorokin@skao.int</a>", self)
        textLabel2.setWordWrap(True)
        textLabel2.setOpenExternalLinks(True)
        textLabel2.setTextFormat(Qt.RichText)
        layout.addWidget(textLabel2, 1, 1)

        self.setLayout(layout)


class InteractiveDatabase(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle('Interactive database window')
        self.parent = parent
        self.country_codes=self.parent.load_country_codes() # type: ignore

        self.desired_width = 1600
        self.desired_height = 900
        self.initUI()

    def initUI(self):
        # Set the window properties. Assuming that we have parent above
        self.setWindowTitle(
            f'Interactive Database Tool connected to database {self.parent.database_version} published on {self.parent.database_date.date()}') # type: ignore
        self.setWindowIcon(QIcon('cps-logo-mono.ico'))

        # Center the window on the screen or parent

        self.parent.centerWindow(self, self.parent) # type: ignore
        self.parent.animateOpening( # type: ignore
            self, self.desired_width, self.desired_height)

        centralWidget = QWidget(self)
        self.setCentralWidget(centralWidget)
        layout = QGridLayout(centralWidget)

        self.tableWidget = QTableWidget()
        headers = ["ITU Notice ID", "ITU Administration code", "ITU Country code        ", "Station Name",
                   "Provision", "Date received", "Longitude", "Latitude"]
        self.tableWidget.setColumnCount(len(headers))
        self.tableWidget.setHorizontalHeaderLabels(headers)
        self.tableWidget.setSortingEnabled(True)
        self.load_data()
        self.tableWidget.cellDoubleClicked.connect(
            self.openDatabaseEntryDetails)
        layout.addWidget(self.tableWidget, 0, 0, 8, 1)

        self.interactionPanel = QGroupBox('Interaction Panel')
        layout_interaction = QGridLayout(self.interactionPanel)
        self.displayNamesCheckbox = QCheckBox(
            "Display ITU Names instead of codes", self.interactionPanel)
        self.displayNamesCheckbox.stateChanged.connect(self.updateTableDisplay)
        checkboxContainer = QWidget()
        checkboxLayout = QHBoxLayout()
        checkboxLayout.setAlignment(Qt.AlignCenter)
        checkboxLayout.addWidget(self.displayNamesCheckbox)
        checkboxContainer.setLayout(checkboxLayout)

        layout_interaction.addWidget(checkboxContainer, 0, 0, 1, 1)
        self.showMapButton = QPushButton(
            "Show stations on the map", self.interactionPanel)
        self.showMapButton.clicked.connect(self.showStationsOnMap)
        layout_interaction.addWidget(self.showMapButton, 1, 0, 1, 1)

        saveCsvButton = QPushButton("Save as CSV", self.interactionPanel)
        layout_interaction.addWidget(saveCsvButton, 0, 1, 1, 1)
        saveCsvButton.clicked.connect(self.saveAsCsv)

        saveWordButton = QPushButton("Save as DOCX", self.interactionPanel)
        layout_interaction.addWidget(saveWordButton, 1, 1, 1, 1)
        saveWordButton.clicked.connect(self.saveAsWord)

        layout.addWidget(self.interactionPanel, 8, 0, 2, 1)

        self.statusBar().showMessage(
            f'Connected to database {self.parent.database_version} published on {self.parent.database_date.date()}') # type: ignore
        self.statusBar().setStyleSheet("""
            QStatusBar {
                background-color: #404040;
                color: white;
                border: 2px solid black;
            }
        """)

        self.updateTableDisplay()
        self.statusBar().showMessage('    Ready')
        self.show()
        self.setFocus()
        self.raise_()
        self.activateWindow()

    def load_data(self):
        SQL = "SELECT ntc_id, adm, ctry, stn_name, prov, d_rcv, long_dec, lat_dec FROM com_el WHERE ntc_type='R' ORDER BY adm asc, stn_name asc;"
        self.rows = self.parent.parse_database(SQL) # type: ignore

        self.tableWidget.setRowCount(len(self.rows))
        for row_num, row_data in enumerate(self.rows):
            for column_num, data in enumerate(row_data):
                if column_num == 5:
                    data = data.strftime("%Y-%m-%d")
                item = QTableWidgetItem(str(data))
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                self.tableWidget.setItem(row_num, column_num, item)

        self.tableWidget.resizeColumnsToContents()

    def updateTableDisplay(self):
        displayNames = self.displayNamesCheckbox.isChecked()
        if displayNames:
            self.tableWidget.horizontalHeaderItem(
                1).setText("ITU Administration name")
            self.tableWidget.horizontalHeaderItem(
                2).setText("ITU Country name")
            self.statusBar().showMessage(
                '    View switched to administration and country name view')
        else:
            self.tableWidget.horizontalHeaderItem(
                1).setText("ITU Administration code")
            self.tableWidget.horizontalHeaderItem(
                2).setText("ITU Country code")
            self.statusBar().showMessage(
                '    View switched to administration and country code view')
        for row_num, row_data in enumerate(self.rows):
            admin_code = row_data[1]
            admin_name = self.country_codes.get(admin_code, "Unknown")
            display_text = admin_name if displayNames else admin_code
            item = QTableWidgetItem(display_text)
            item.setFlags(item.flags() & ~Qt.ItemIsEditable)

            tooltip_text = admin_name if not displayNames else None
            item.setToolTip(tooltip_text)
            self.tableWidget.setItem(row_num, 1, item)

            country_code = row_data[2]
            country_name = self.country_codes.get(country_code, "Unknown")
            display_text = country_name if displayNames else country_code
            item = QTableWidgetItem(display_text)
            item.setFlags(item.flags() & ~Qt.ItemIsEditable)

            tooltip_text = country_name if not displayNames else None
            item.setToolTip(tooltip_text)

            self.tableWidget.setItem(row_num, 2, item)

    def showStationsOnMap(self):
        station_data = []

        for row in range(self.tableWidget.rowCount()):
            raw_adm_info = self.tableWidget.item(row, 1).text()
            raw_country_info = self.tableWidget.item(row, 2).text()
            station_name = self.tableWidget.item(row, 3).text()
            longitude = self.tableWidget.item(row, 6).text()
            latitude = self.tableWidget.item(row, 7).text()

            if self.displayNamesCheckbox.isChecked():
                code = [key for key, value in self.country_codes.items()
                        if value == raw_adm_info]
                administration_info = f"{raw_adm_info} ({code[0]})" if code else raw_adm_info
                code = [key for key, value in self.country_codes.items()
                        if value == raw_country_info]
                country_info = f"{raw_country_info} ({code[0]})" if code else raw_country_info
            else:
                administration_name = self.country_codes.get(
                    raw_adm_info, "Unknown Country")
                administration_info = f"{administration_name} ({raw_adm_info})"
                country_name = self.country_codes.get(
                    raw_country_info, "Unknown Country")
                country_info = f"{country_name} ({raw_country_info})"

            try:
                station_data.append((station_name, administration_info, country_info, float(
                    latitude), float(longitude)))
            except:
                station_data.append((station_name, administration_info, country_info, (
                    latitude), (longitude)))

        self.parent.animateClosing(self) # type: ignore
        self.showMinimized()
        self.setEnabled(False)
        self.mapWindow = MapWindow(station_data, self)

    def openDatabaseEntryDetails(self, row, column):
        ntc_id_item = self.tableWidget.item(row, 0)
        station_name_item = self.tableWidget.item(row, 3)

        ntc_id = ntc_id_item.text() if ntc_id_item else "Unknown"
        station_name = station_name_item.text() if station_name_item else "Unknown"

        self.parent.animateClosing(self) # type: ignore
        self.showMinimized()
        self.setEnabled(False)
        self.detailsWindow = DatabaseEntryDetails(ntc_id, station_name, self)

    def saveAsCsv(self):
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save as CSV", f"RAS_DB_OVERVIEW_CSV_{self.parent.database_version}_{self.parent.database_date.date()}", "CSV Files (*.csv)") # type: ignore
        if filePath:
            if os.path.basename(filePath) == 'geographical-areas.csv':
                QMessageBox.critical(
                    self, "Error", "This file is an important app file and cannot be overwritten.")
                return
            try:
                with open(filePath, 'w', newline='', encoding='utf-8-sig') as file:
                    writer = csv.writer(file)
                    headers = [self.tableWidget.horizontalHeaderItem(
                        i).text() for i in range(self.tableWidget.columnCount())]
                    writer.writerow(headers)
                    for row in range(self.tableWidget.rowCount()):
                        rowData = [self.tableWidget.item(row, i).text() if self.tableWidget.item(
                            row, i) else '' for i in range(self.tableWidget.columnCount())]
                        writer.writerow(rowData)
            except Exception as e:
                QMessageBox.critical(self, "CSV saving Error",
                                     f"An error occurred while preparing the csv file:\n{e}")

    def saveAsWord(self):
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save as DOCX", f"RAS_DB_OVERVIEW_DOCX_{self.parent.database_version}_{self.parent.database_date.date()}", "Word Files (*.docx)") # type: ignore
        if filePath:
            doc = docx.Document()
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

            doc.add_heading(
                f'Overview RAS information as per database {self.parent.database_version} published on {self.parent.database_date.date()}', level=1) # type: ignore

            table = doc.add_table(rows=1, cols=self.tableWidget.columnCount())
            header_cells = table.rows[0].cells

            for column in range(self.tableWidget.columnCount()):
                header = self.tableWidget.horizontalHeaderItem(column)
                if header:
                    header_cells[column].text = header.text()

            for row in range(self.tableWidget.rowCount()):
                row_cells = table.add_row().cells
                for col in range(self.tableWidget.columnCount()):
                    item = self.tableWidget.item(row, col)
                    row_cells[col].text = item.text() if item else ""

            try:
                doc.save(filePath)
            except Exception as e:
                QMessageBox.critical(self, "Docx saving Error",
                                     f"An error occurred while preparing the docx file:\n{e}")

    def closeEvent(self, event):
        if self.parent:
            self.parent.animateClosing(self)
            self.parent.showNormal()
            self.parent.animateOpening(
                self.parent, self.parent.desired_width, self.parent.desired_height)
            self.parent.setFocus()
            self.parent.raise_()
            self.parent.activateWindow()
            self.parent.setEnabled(True)

class LoadingWidget(QWidget):
    def __init__(self):
        super().__init__()
        layout = QGridLayout(self)
        layout.setAlignment(Qt.AlignCenter)

        self.label = QLabel("Map is being processed, please wait")
        self.label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.label, 0, 0, Qt.AlignCenter)

        self.spinner = SpinnerWidget()
        layout.addWidget(self.spinner, 1, 0, Qt.AlignCenter)

        self.setLayout(layout)


class SpinnerWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.setFixedSize(60, 60)
        self.angle = 0
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.rotate)
        self.timer.start(100)

    def rotate(self):
        self.angle = (self.angle + 30) % 360
        self.update()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.setPen(Qt.NoPen)

        rect = self.rect()
        radius = min(rect.width(), rect.height()) // 2

        for i in range(12):
            color = QColor(0, 0, 0)
            color.setAlphaF(1.0 - i / 12.0)
            painter.setBrush(color)

            painter.save()
            painter.translate(rect.center())
            painter.rotate(self.angle - i * 30.0)
            painter.drawEllipse(QRectF(-radius // 4, -radius, radius // 2, radius // 2))
            painter.restore()

class MapWindow(QMainWindow):
    def __init__(self, station_data=None, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.station_data = station_data
        self.initUI()

    def initUI(self):
        self.setWindowTitle(
            f'Radio astronomy stations on a map as per database {self.parent.parent.database_version} published on {self.parent.parent.database_date.date()}') # type: ignore
        self.setWindowIcon(QIcon('cps-logo-mono.ico'))

        self.parent.parent.centerWindow(self, self.parent) # type: ignore
        self.parent.parent.animateOpening( # type: ignore
            self, desired_width=1600, desired_height=900)

        centralWidget = QWidget()
        self.setCentralWidget(centralWidget)
        layout = QGridLayout()
        centralWidget.setLayout(layout)

        self.mapPanel = QGroupBox("Map Panel")
        self.mapLayout = QStackedLayout()
        self.mapPanel.setLayout(self.mapLayout)

        self.browser = QWebEngineView()
        self.loading_widget = LoadingWidget()

        self.mapLayout.addWidget(self.browser)
        self.mapLayout.addWidget(self.loading_widget)
        self.mapLayout.setCurrentWidget(self.loading_widget)

        self.browser.loadFinished.connect(self.onLoadFinished)

        self.interactionPanel = QGroupBox("Interaction Panel")
        interactionLayout = QGridLayout()
        self.interactionPanel.setLayout(interactionLayout)

        saveScreenshotButton = QPushButton(
            "Save as Screenshot", self.interactionPanel)
        interactionLayout.addWidget(saveScreenshotButton, 0, 0)
        saveScreenshotButton.clicked.connect(self.saveScreenshot)

        saveHTMLButton = QPushButton("Save as HTML", self.interactionPanel)
        interactionLayout.addWidget(saveHTMLButton, 1, 0)
        saveHTMLButton.clicked.connect(self.saveHTML)

        layout.addWidget(self.mapPanel, 0, 0)
        layout.addWidget(self.interactionPanel, 1, 0)
        layout.setRowStretch(0, 8)
        layout.setRowStretch(1, 2)

        self.show()
        self.setFocus()
        self.raise_()
        self.activateWindow()

        if self.station_data is not None:
            self.browser.setHtml(self.generateMapHTML(self.station_data))

    def onLoadFinished(self, ok):
        if ok:
            self.mapLayout.setCurrentWidget(self.browser)

    def generate_base64_icon(self, icon_path):
        with open(icon_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
        
    def generateMapHTML(self, station_data):
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            icon_path = os.path.join(script_dir, 'ras_s_icon.webp').replace('\\', '/')
            icon_base64 = self.generate_base64_icon(icon_path)

            html_parts = []
            html_parts.append(f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Full Widget Leaflet Map for Radio Astronomy Stations Database Tool</title>
                <meta charset="utf-8" />
                <link 
                    rel="stylesheet" 
                    href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"
                />
                <link
                    rel="stylesheet"
                    href="https://unpkg.com/leaflet.markercluster/dist/MarkerCluster.css"
                />
                <link
                    rel="stylesheet"
                    href="https://unpkg.com/leaflet.markercluster/dist/MarkerCluster.Default.css"
                />
                <script 
                    src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js">
                </script>
                <script
                    src="https://unpkg.com/leaflet.markercluster/dist/leaflet.markercluster.js">
                </script>
                <style>
                    body {{
                        padding: 0;
                        margin: 0;
                    }}
                    html, body, #map {{
                        height: 100%;
                        width: 100%;
                    }}
                    .custom-cluster-icon {{
                        background: radial-gradient(circle, white 25%, transparent 75%);
                        border-radius: 50%;
                        border: 2px solid rgba(0, 0, 0, 0.5);
                        text-align: center;
                        color: black;
                        font-size: 14px;
                        font-weight: bold;
                        width: 40px;
                        height: 40px;
                    }}
                    .custom-cluster-icon img {{
                        position: absolute;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%);
                        width: 30px;
                        height: 30px;
                    }}
                    .custom-cluster-icon .cluster-count {{
                        position: absolute;
                        top: -10px;
                        right: -10px;
                        background: red;
                        color: white;
                        border-radius: 50%;
                        width: 20px;
                        height: 20px;
                        display: flex;
                        justify-content: center;
                        align-items: center;
                        font-size: 12px;
                    }}
                </style>
            </head>
            <body>
                <div id="map"></div>
            
                <script>
                    var map = L.map('map', {{attributionControl: false}}).setView([0, 0], 2);
                    var myAttrControl = L.control.attribution().addTo(map);
                    myAttrControl.setPrefix('<a href="https://leafletjs.com/">Leaflet</a>');
                    
                    mapLink = 
                        '<a href="http://openstreetmap.org">OpenStreetMap</a>';
                    L.tileLayer(
                        'http://{{s}}.tile.openstreetmap.org/{{z}}/{{x}}/{{y}}.png', {{
                        attribution: 'Map data by &copy; <a href="http://www.openstreetmap.org/copyright">OpenStreetMap</a>, under <a href="https://opendatacommons.org/licenses/odbl/">ODbL.</a>',
                        maxZoom: 18,
                        }}).addTo(map);
                    
                    var customIcon = L.icon({{
                        iconUrl: 'data:image/webp;base64,{icon_base64}',
                        iconSize: [30, 30],
                        iconAnchor: [15, 15],
                        popupAnchor: [0, 0]
                    }}); 

                    var markers = L.markerClusterGroup({{
                        maxClusterRadius: 50,
                        iconCreateFunction: function(cluster) {{
                            var childCount = cluster.getChildCount();
                            return L.divIcon({{
                                html: '<div><img src="data:image/webp;base64,{icon_base64}" alt="cluster-icon"/><div class="cluster-count">' + childCount + '</div></div>',
                                className: 'custom-cluster-icon',
                                iconSize: [40, 40]
                            }});
                        }}
                    }});
            """)

            for name, adm, ctr, lat, lon in station_data:
                adm_escaped = adm.replace("'", "&#39;").replace('"', '&quot;')
                ctr_escaped = ctr.replace("'", "&#39;").replace('"', '&quot;')
                html_parts.append(f"""                                  
                                var marker = L.marker([{lat}, {lon}], {{icon: customIcon}}).bindPopup('<b>{name}</b><br>Region country code:<b>{ctr_escaped}</b><br>Responsible administration: <b>{adm_escaped}</b>');
                                markers.addLayer(marker);
                                """)

            html_parts.append("""
                    map.addLayer(markers);
                </script>
            </body>
            </html>
            """)

            html = ''.join(html_parts)

        except Exception as e:
            html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Error</title>
        </head>
        <body>
            <p>Error generating map: {str(e)}</p>
        </body>
        </html>
        """
        return html

    def saveScreenshot(self):
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save Screenshot", f"RAS_DB_MAP_{self.parent.parent.database_version}_{self.parent.parent.database_date.date()}", "PNG Files (*.png)") # type: ignore
        if filePath:
            self.browser.grab().save(filePath)

    def saveHTML(self):
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save HTML", f"RAS_DB_MAP_{self.parent.parent.database_version}_{self.parent.parent.database_date.date()}", "HTML Files (*.html)") # type: ignore
        if filePath:
            def save_html(html):
                with open(filePath, "w", encoding='utf-8') as file:
                    file.write(html)
            self.browser.page().toHtml(save_html)

    def closeEvent(self, event):
        if self.parent:
            self.parent.parent.animateClosing(self)
            self.parent.showNormal()
            self.parent.parent.animateOpening(
                self.parent, self.parent.desired_width, self.parent.desired_height)
            self.parent.setFocus()
            self.parent.raise_()
            self.parent.activateWindow()
            self.parent.setEnabled(True)


class DatabaseEntryDetails(QMainWindow):
    def __init__(self, ntc_id=None, station_name=None, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.ntc_id = ntc_id
        self.station_name = station_name
        self.station_rows = []
        self.initUI()

    def initUI(self):
        self.setWindowTitle(
            f'Details of station {self.station_name} with notice id {self.ntc_id}')
        self.setWindowIcon(QIcon('cps-logo-mono.ico'))

        # Center the window on the screen or parent

        self.parent.parent.centerWindow(self, self.parent) # type: ignore
        self.parent.parent.animateOpening( # type: ignore
            self, desired_width=1600, desired_height=900)

        centralWidget = QWidget(self)
        self.setCentralWidget(centralWidget)
        layout = QGridLayout(centralWidget)

        stationInfoPanel = QGroupBox('Station Information')
        self.stationInfoLayout = QGridLayout(stationInfoPanel)
        self.stationInfoTable = QTableWidget()
        headers = ['Notice ID', 'Longitude (Degrees)', 'East/West', 'Minutes', 'Seconds', 'Latitude (Degrees)',
                   'North/South', 'Minutes', 'Seconds', 'Min elevation', 'Max elevation', 'Min azimuth', 'Max Azimuth',
                   'Antenna altitude, m']
        self.stationInfoTable.setColumnCount(len(headers))
        self.stationInfoTable.setHorizontalHeaderLabels(headers)
        self.stationInfoTable.setSortingEnabled(True)

        self.stationInfoLayout.addWidget(self.stationInfoTable)
        layout.addWidget(stationInfoPanel, 0, 0, 1, 4)

        self.beamInfoPanel = QGroupBox('Beams Information')
        self.beamInfoLayout = QGridLayout(self.beamInfoPanel)
        self.beamInfoTable = QTableWidget()
        headers = ['Beam name', 'Antenna Code', 'Antenna diameter, m', 'Antenna gain, dBi',
                   'Noise temp, K', 'Frequency minimum, MHz', 'Frequency maximum, MHz', 'VLBI type', 'Centre frequency, MHz']
        self.beamInfoTable.setColumnCount(len(headers))
        self.beamInfoTable.setHorizontalHeaderLabels(headers)
        self.beamInfoTable.setSortingEnabled(True)

        self.beamInfoLayout.addWidget(self.beamInfoTable)
        layout.addWidget(self.beamInfoPanel, 1, 0, 1, 4)

        self.controlPanel = QGroupBox('Interaction Panel')
        self.controlLayout = QGridLayout(self.controlPanel)
        layout.addWidget(self.controlPanel, 2, 0, 1, 1)

        openAntennaTableButton = QPushButton(
            "Open Antenna Code Table (System PDF viewer)", self.controlPanel)
        self.controlLayout.addWidget(openAntennaTableButton, 0, 0)
        openAntennaTableButton.clicked.connect(self.openAntennaTable)

        saveCsvButton = QPushButton("Save Tables to CSV", self.controlPanel)
        self.controlLayout.addWidget(saveCsvButton)
        saveCsvButton.clicked.connect(self.saveTablesToCsv)

        saveDocxButton = QPushButton("Save Tables to DOCX", self.controlPanel)
        self.controlLayout.addWidget(saveDocxButton)
        saveDocxButton.clicked.connect(self.saveTablesToWord)

        self.mapPanel = QGroupBox('Map Panel')
        self.mapLayout = QGridLayout(self.mapPanel)

        self.browser = QWebEngineView()

        layout.addWidget(self.mapPanel, 2, 1, 1, 3)

        self.mapLayout.addWidget(self.browser, 0, 0)

        self.load_data()
        self.browser.setHtml(self.generateMapHTML())

        layout.setRowStretch(0, 1)
        layout.setRowStretch(1, 4)
        layout.setRowStretch(2, 5)

        self.show()
        self.setFocus()
        self.raise_()
        self.activateWindow()

    def load_data(self):
        SQL = f"SELECT long_deg, long_ew, long_min, long_sec, lat_deg, lat_ns, lat_min, lat_sec, elev_min, elev_max, azm_fr, azm_to, ant_alt FROM e_stn WHERE ntc_id={str(self.ntc_id)};"
        self.station_rows = self.parent.parent.parse_database(SQL) # type: ignore

        self.stationInfoTable.setRowCount(len(self.station_rows))
        for row_num, row_data in enumerate(self.station_rows):
            item = QTableWidgetItem(str(self.ntc_id))
            item.setFlags(item.flags() & ~Qt.ItemIsEditable)
            self.stationInfoTable.setItem(row_num, 0, item)

            for column_num, data in enumerate(row_data):
                if data == None:
                    data = 'N/A'
                item = QTableWidgetItem(str(data))
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                self.stationInfoTable.setItem(row_num, column_num+1, item)

        self.stationInfoTable.resizeColumnsToContents()

        SQL = f"SELECT beam_name, pattern_id, ant_diam, gain FROM e_ant WHERE ntc_id={str(self.ntc_id)};"
        self.beam_rows = self.parent.parent.parse_database(SQL) # type: ignore

        self.beamInfoTable.setRowCount(0)
        for beam_index, beam_row_data in enumerate(self.beam_rows):
            beam_name = beam_row_data[0]
            SQL = f"SELECT noise_t, freq_min, freq_max, ra_stn_type FROM grp WHERE ntc_id={str(self.ntc_id)} AND beam_name='{beam_name}';"
            self.grp_rows = self.parent.parent.parse_database(SQL) # type: ignore

            SQL = f"SELECT freq_mhz FROM freq WHERE ntc_id={str(self.ntc_id)} AND beam_name='{beam_name}';"
            self.freq_rows = self.parent.parent.parse_database(SQL) # type: ignore

            if not beam_row_data[1] == None:
                SQL = f"SELECT pattern FROM ant_type WHERE pattern_id={beam_row_data[1]};"
                try:
                    antenna_code = self.parent.parent.parse_database(SQL)[0][0] # type: ignore
                except:
                    try:
                        antenna_code = self.parent.parent.parse_database(SQL)[ # type: ignore
                            0]
                    except:
                        antenna_code = None
                if antenna_code == None:
                    antenna_code = 'N/A'
            else:
                antenna_code = 'N/A'

            for grp_index, grp_row_data in enumerate(self.grp_rows):
                current_row_count = self.beamInfoTable.rowCount()
                self.beamInfoTable.insertRow(current_row_count)
                freq_row_data = self.freq_rows[grp_index][0]

                combined_row_data = list(
                    beam_row_data) + list(grp_row_data) + [freq_row_data]

                for column_num, data in enumerate(combined_row_data):
                    if data == None:
                        data = 'N/A'
                    if column_num == 1:
                        data = antenna_code
                    if column_num == 7:
                        if data == 'S':
                            data = 'Single'
                        elif data == 'V':
                            data = 'VLBI'
                    item = QTableWidgetItem(str(data))
                    item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                    self.beamInfoTable.setItem(
                        current_row_count, column_num, item)

        self.beamInfoTable.resizeColumnsToContents()

    def generateMapHTML(self):
        SQL = f"SELECT adm, ctry, stn_name, long_dec, lat_dec FROM com_el WHERE ntc_id={str(self.ntc_id)};"
        self.rows = self.parent.parent.parse_database(SQL)[0] # type: ignore

        adm = self.rows[0]
        ctr = self.rows[1]

        administration_name = self.parent.country_codes.get( # type: ignore
            adm, "Unknown Country")
        adm = f"{administration_name} ({adm})"
        country_name = self.parent.country_codes.get(ctr, "Unknown Country") # type: ignore
        ctr = f"{country_name} ({ctr})"

        name = self.rows[2]
        lon = float(self.rows[3])
        lat = float(self.rows[4])

        try:
            html_parts = []
            html_parts.append("""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Full Widget Leaflet Map for Radio Astronomy Stations Database Tool</title>
                <meta charset="utf-8" />
                <link 
                    rel="stylesheet" 
                    href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"
                />
                <script 
                    src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js">
                </script>
                <style>
                    body {
                        padding: 0;
                        margin: 0;
                    }
                    html, body, #map {
                        height: 100%;
                        width: 100%;
                    }
                </style>
            </head>
            <body>
                <div id="map"></div>
            
                <script>
                """)
            html_parts.append(f"""
                    var map = L.map('map', {{attributionControl: false}}).setView([{lat}, {lon}], 5);
                    """)
            html_parts.append("""
                    var myAttrControl = L.control.attribution().addTo(map);
                    myAttrControl.setPrefix('<a href="https://leafletjs.com/">Leaflet</a>');
                    
                    mapLink = 
                        '<a href="http://openstreetmap.org">OpenStreetMap</a>';
                    L.tileLayer(
                        'http://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png', {
                        attribution: 'Map data by &copy; <a href="http://www.openstreetmap.org/copyright">OpenStreetMap</a>, under <a href="https://opendatacommons.org/licenses/odbl/">ODbL.</a>',
                        maxZoom: 18,
                        }).addTo(map);
            """)
            adm_escaped = adm.replace("'", "&#39;").replace('"', '&quot;')
            ctr_escaped = ctr.replace("'", "&#39;").replace('"', '&quot;')
            html_parts.append(
                f"        L.marker([{lat}, {lon}]).addTo(map).bindPopup('<b>{name}</b><br>Region country code:<b>{ctr_escaped}</b><br>Responsible administration: <b>{adm_escaped}</b>').openPopup();")
            html_parts.append("""
                </script>
            </body>
            </html>
            """)

            html = ''.join(html_parts)

        except Exception as e:
            html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Error</title>
        </head>
        <body>
            <p>Error generating map: {str(e)}</p>
        </body>
        </html>
        """
        return html

    def openAntennaTable(self):
        QDesktopServices.openUrl(QUrl.fromLocalFile('Table6toPreface.pdf'))

    def saveTablesToCsv(self):
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save as CSV", f"RAS_DB_CSV_{self.station_name}_{self.parent.parent.database_version}_{self.parent.parent.database_date.date()}", "CSV Files (*.csv)") # type: ignore
        if filePath:
            if os.path.basename(filePath) == 'geographical-areas.csv':
                QMessageBox.critical(
                    self, "Error", "This file is an important app file and cannot be overwritten.")
                return
            try:
                with open(filePath, 'w', newline='', encoding='utf-8-sig') as file:
                    writer = csv.writer(file)

                    headers = [self.stationInfoTable.horizontalHeaderItem(
                        i).text() for i in range(self.stationInfoTable.columnCount())]
                    writer.writerow(headers)
                    for row in range(self.stationInfoTable.rowCount()):
                        rowData = [self.stationInfoTable.item(row, i).text() if self.stationInfoTable.item(
                            row, i) else '' for i in range(self.stationInfoTable.columnCount())]
                        writer.writerow(rowData)

                    writer.writerow([])

                    headers = [self.beamInfoTable.horizontalHeaderItem(
                        i).text() for i in range(self.beamInfoTable.columnCount())]
                    writer.writerow(headers)
                    for row in range(self.beamInfoTable.rowCount()):
                        rowData = [self.beamInfoTable.item(row, i).text() if self.beamInfoTable.item(
                            row, i) else '' for i in range(self.beamInfoTable.columnCount())]
                        writer.writerow(rowData)
            except Exception as e:
                QMessageBox.critical(self, "CSV saving Error",
                                     f"An error occurred while preparing the csv file:\n{e}")

    def saveTablesToWord(self):
        filePath, _ = QFileDialog.getSaveFileName(
            self, "Save as DOCX", f"RAS_DB_DOCX_{self.station_name}_{self.parent.parent.database_version}_{self.parent.parent.database_date.date()}", "Word Files (*.docx)") # type: ignore
        if filePath:
            doc = docx.Document()
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

            doc.add_heading(
                f'Station "{self.station_name}" information as per database {self.parent.parent.database_version} published on {self.parent.parent.database_date.date()}', level=1) # type: ignore
            doc.add_heading("General information", level=2)
            SQL = f"SELECT adm, ctry, long_dec, lat_dec FROM com_el WHERE ntc_id={self.ntc_id};"
            rows = self.parent.parent.parse_database(SQL)[0] # type: ignore
            adm = rows[0]
            administration_name = self.parent.country_codes.get( # type: ignore
                adm, "Unknown Country")
            doc.add_paragraph(
                f"Responsible administration: {administration_name}")
            ctry = rows[1]
            country_name = self.parent.country_codes.get( # type: ignore
                ctry, "Unknown Country")
            doc.add_paragraph(f"Region country code: {country_name}")
            long_dec = rows[2]
            doc.add_paragraph(
                f"Longitude {abs(long_dec)} {self.station_rows[0][1]}")
            lat_dec = rows[3]
            doc.add_paragraph(
                f"Latitude {abs(lat_dec)} {self.station_rows[0][5]}")

            doc.add_heading("Station Information Table", level=2)

            doc.add_paragraph("")
            table = doc.add_table(
                rows=1, cols=self.stationInfoTable.columnCount())
            header_cells = table.rows[0].cells

            for col in range(self.stationInfoTable.columnCount()):
                header_cells[col].text = self.stationInfoTable.horizontalHeaderItem(
                    col).text()
            for row in range(self.stationInfoTable.rowCount()):
                row_cells = table.add_row().cells
                for col in range(self.stationInfoTable.columnCount()):
                    row_cells[col].text = self.stationInfoTable.item(
                        row, col).text() if self.stationInfoTable.item(row, col) else ''

            doc.add_paragraph("")

            doc.add_heading("Beams Information Table", level=2)

            doc.add_paragraph("")

            table = doc.add_table(
                rows=1, cols=self.beamInfoTable.columnCount())
            header_cells = table.rows[0].cells

            for col in range(self.beamInfoTable.columnCount()):
                header_cells[col].text = self.beamInfoTable.horizontalHeaderItem(
                    col).text()
            for row in range(self.beamInfoTable.rowCount()):
                row_cells = table.add_row().cells
                for col in range(self.beamInfoTable.columnCount()):
                    row_cells[col].text = self.beamInfoTable.item(
                        row, col).text() if self.beamInfoTable.item(row, col) else ''

            doc.add_paragraph("")

            try:
                doc.save(filePath)
            except Exception as e:
                QMessageBox.critical(self, "Docx saving Error",
                                     f"An error occurred while preparing the docx file:\n{e}")

    def closeEvent(self, event):
        if self.parent:
            self.parent.parent.animateClosing(self)
            self.parent.showNormal()
            self.parent.parent.animateOpening(
                self.parent, self.parent.desired_width, self.parent.desired_height)
            self.parent.setFocus()
            self.parent.raise_()
            self.parent.activateWindow()
            self.parent.setEnabled(True)

class SiteLinkWizard(QMainWindow):
    def __init__(self, filePath=None, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.filePath = filePath
        self.country_codes=self.parent.load_country_codes() # type: ignore
        self.desired_width=1600
        self.desired_height=900
        self.map_html = ""
        self.initUI()

    def initUI(self):
        self.setWindowTitle(
            f'Site Link Wizard - A magical way to link Wikidata entries to ITU DB entries')
        self.setWindowIcon(QIcon('cps-logo-mono.ico'))
        self.parent.centerWindow(self, self.parent) # type: ignore
        self.parent.animateOpening( # type: ignore
            self, desired_width=self.desired_width, desired_height=self.desired_height)
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        self.layout = QGridLayout(central_widget)

        self.instructions = QLabel("Find and link corresponding stations.")
        self.layout.addWidget(self.instructions, 0, 0, 1, 4)

        self.progressBar = QProgressBar()
        self.layout.addWidget(self.progressBar, 1, 0, 1, 4)

        self.wikidataGroup = QGroupBox("Wikidata Entry")
        self.wikidataLayout = QGridLayout()
        self.wikidataGroup.setLayout(self.wikidataLayout)
        self.layout.addWidget(self.wikidataGroup, 2, 0, 2, 1)

        self.stationsGroup = QGroupBox("Closest Stations registered at ITU")
        self.stationsLayout = QGridLayout()
        self.stationsGroup.setLayout(self.stationsLayout)
        self.layout.addWidget(self.stationsGroup, 2, 1, 2, 1)

        self.stationsList = QListWidget()
        self.stationsList.itemSelectionChanged.connect(self.display_station_details)
        self.stationsLayout.addWidget(self.stationsList, 0, 0)

        self.stationDetailsGroup = QGroupBox("ITU Entry Details")
        self.stationDetailsLayout = QGridLayout()
        self.stationDetailsGroup.setLayout(self.stationDetailsLayout)
        self.layout.addWidget(self.stationDetailsGroup, 2, 2, 1, 2)

        self.showDetailsButton = QPushButton("Show ITU Entry Details")
        self.showDetailsButton.clicked.connect(self.show_itu_details)
        self.layout.addWidget(self.showDetailsButton, 3, 2, 1, 2)

        self.MapGroup = QGroupBox("Station locations")
        self.mapLayout = QGridLayout()
        self.MapGroup.setLayout(self.mapLayout)
        self.browser = QWebEngineView()
        self.mapLayout.addWidget(self.browser)
        self.layout.addWidget(self.MapGroup, 4, 0, 2, 4)

        self.buttonsLayout = QHBoxLayout()
        self.confirmButton = QPushButton("Confirm Match")
        self.noMatchButton = QPushButton("No Match")

        self.buttonsLayout.addWidget(self.confirmButton)
        self.buttonsLayout.addWidget(self.noMatchButton)

        self.layout.addLayout(self.buttonsLayout, 6, 0, 1, 4)

        self.confirmButton.clicked.connect(self.confirm_match)
        self.noMatchButton.clicked.connect(self.no_match)

        
        self.load_data()
        self.current_index = 0
        self.show_entry()

        self.show()
        self.setFocus()
        self.raise_()
        self.activateWindow()

    def load_data(self):
        self.conn = sqlite3.connect(self.filePath) # type: ignore
        self.cursor = self.conn.cursor()

        # Drop link table if already exist
        drop_link_table_query = "DROP TABLE IF EXISTS wikidata_stations_link"
        self.cursor.execute(drop_link_table_query)

        # Create the link table
        create_link_table_query = """
        CREATE TABLE wikidata_stations_link (
            "CPS Wiki ID" INTEGER,
            "CPS Station ID" INTEGER,
            FOREIGN KEY("CPS Wiki ID") REFERENCES "wikidata"("CPS Wiki ID"),
            FOREIGN KEY("CPS Station ID") REFERENCES "Stations"("CPS Station ID")
        );
        """
        self.cursor.execute(create_link_table_query)
        
        self.cursor.execute("SELECT * FROM wikidata")
        self.wikidata_entries = self.cursor.fetchall()

        self.cursor.execute("SELECT * FROM Stations")
        self.stations_entries = self.cursor.fetchall()

        self.progressBar.setMaximum(len(self.wikidata_entries))
        self.progressBar.setValue(0)

        # Convert stations entries to numpy arrays for efficient calculations
        self.stations_coords = np.array([(station[6], station[5]) for station in self.stations_entries])
        self.station_names = [station[2] for station in self.stations_entries]

    def show_entry(self):
        if self.current_index < len(self.wikidata_entries):
            self.station_data = []
            entry = self.wikidata_entries[self.current_index]
            self.display_wikidata_entry(entry)
            self.find_closest_stations(entry)
            self.progressBar.setValue(self.current_index + 1)
            
            latitude = entry[4]
            longitude = entry[3]
            if latitude is None or longitude is None:
                self.confirmButton.setEnabled(False)
            else:
                self.confirmButton.setEnabled(True)

            self.map_html = self.generateMapHTML(self.station_data)
            self.browser.setHtml(self.map_html)
        else:
            self.instructions.setText("No more entries.")
            self.confirmButton.setEnabled(False)
            self.noMatchButton.setEnabled(False)
            self.clear_layout(self.wikidataLayout)
            self.clear_layout(self.stationDetailsLayout)

    def clear_layout(self, layout):
        for i in reversed(range(layout.count())):
            item = layout.itemAt(i)
            if item.widget() is not None:
                item.widget().deleteLater()
            elif item.layout() is not None:
                self.clear_layout(item.layout())
            layout.removeItem(item)

    def display_wikidata_entry(self, entry):
        # Clear the previous entries
        self.clear_layout(self.wikidataLayout)

        self.wikidataLayout.addWidget(QLabel(f"Name: {entry[1]}"), 0, 0)
        self.wikidataLayout.addWidget(QLabel(f"Country: {entry[2]}"), 1, 0)
        self.wikidataLayout.addWidget(QLabel(f"Coordinates:"), 2, 0)
        try:
            self.wikidataLayout.addWidget(QLabel(f"    Longitude: {entry[3]:.3f}"), 3, 0)
            self.wikidataLayout.addWidget(QLabel(f"    Latitude: {entry[4]:.3f}"), 4, 0)
        except:
            self.wikidataLayout.addWidget(QLabel(f"    Longitude: N/A"), 3, 0)
            self.wikidataLayout.addWidget(QLabel(f"    Latitude: N/A"), 4, 0)

        source_label = QLabel(f'<a href="{entry[5]}">Source</a>')
        source_label.setOpenExternalLinks(True)
        self.wikidataLayout.addWidget(source_label, 5, 0)

        # Add a vertical spacer to push the content to the top
        self.wikidataLayout.addItem(QSpacerItem(20, 40, QSizePolicy.Minimum, QSizePolicy.Expanding), 6, 0)
        self.station_data.append((entry[1], "WikiData", entry[2], entry[4], entry[3]))

    def display_station_details(self):
        self.clear_layout(self.stationDetailsLayout)
        selected_items = self.stationsList.selectedItems()
        if selected_items:
            selected_item = selected_items[0].text()
            station_name = selected_item.split(' (')[0]
            matched_station = next((s for s in self.stations_entries if s[2] == station_name), None)
            if matched_station:
                self.stationDetailsLayout.setSpacing(2)
                self.stationDetailsLayout.setContentsMargins(0, 0, 0, 0)

                self.stationDetailsLayout.addWidget(QLabel(f"Name: {matched_station[2]}\n"), 0, 0)
                self.stationDetailsLayout.addWidget(QLabel(f"Country: {matched_station[1]}\n"), 1, 0)
                self.stationDetailsLayout.addWidget(QLabel(f"Coordinates: {matched_station[5]:.3f}, {matched_station[6]:.3f}\n"), 2, 0)
                self.stationDetailsLayout.addWidget(QLabel(f"Frequency range: {matched_station[10]} MHz - {matched_station[11]} MHz\n"), 3, 0)
                self.stationDetailsLayout.addWidget(QLabel(f"ITU Notice ID: {matched_station[17]}\n"), 4, 0)

                # Add a vertical spacer to push the content to the top
                self.stationDetailsLayout.addItem(QSpacerItem(20, 40, QSizePolicy.Minimum, QSizePolicy.Expanding), 5, 0)        

    def show_itu_details(self):
        if self.parent.dbConnection is None: # type: ignore
            QMessageBox.warning(self, "Warning", "ITU database was not connected, please select one corresponding to the creation of this CPS database.")
            self.parent.database_select() # type: ignore
            if self.parent.database_file_name: # type: ignore
                self.parent.database_connect() # type: ignore
            else:
                QMessageBox.warning(self, "Warning", "ITU database selection cancelled")

        if self.parent.dbConnection: # type: ignore
            selected_items = self.stationsList.selectedItems()
            if selected_items:
                station_name = selected_items[0].text().split(' (')[0]
                matched_station = next((s for s in self.stations_entries if s[2] == station_name), None)
                if matched_station:
                    self.parent.animateClosing(self) # type: ignore
                    self.showMinimized()
                    self.setEnabled(False)
                    self.details_window = DatabaseEntryDetails(ntc_id=str(matched_station[17]), station_name=matched_station[2], parent=self)
        else:                    
            QMessageBox.critical(self, "Database Connection Error",
                                    "An error occurred while connecting to the database, no details will be shown")


        

    def find_closest_stations(self, entry):
        latitude = entry[4]
        longitude = entry[3]

        if latitude is not None and longitude is not None:
            distances = self.calculate_distances(latitude, longitude)
            # Select 10 closest stations
            closest_indices = np.argsort(distances)[:10]

            self.stationsList.clear()
            for idx in closest_indices:
                station_name = self.station_names[idx]
                distance = distances[idx]
                item = QListWidgetItem(f"{station_name} ({distance:.2f} km)")
                item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                item.setCheckState(Qt.Unchecked)
                self.stationsList.addItem(item)

                matched_station = next((s for s in self.stations_entries if s[2] == station_name), None)
                if matched_station:
                    self.station_data.append((matched_station[2], matched_station[18], matched_station[1], matched_station[6], matched_station[5]))

            self.stationsList.setCurrentRow(0)
        else:
            self.stationsList.clear()
            self.stationsList.addItem("No coordinates available.")

    def calculate_distances(self, lat1, lon1):
        lat1_rad, lon1_rad = np.radians(lat1), np.radians(lon1)
        lat2_rad, lon2_rad = np.radians(self.stations_coords[:, 0]), np.radians(self.stations_coords[:, 1])
        
        dlat = lat2_rad - lat1_rad
        dlon = lon2_rad - lon1_rad
        
        a = np.sin(dlat / 2)**2 + np.cos(lat1_rad) * np.cos(lat2_rad) * np.sin(dlon / 2)**2
        c = 2 * np.arctan2(np.sqrt(a), np.sqrt(1 - a))
        distance = 6371 * c  # Earth radius in kilometers
        return distance
    
    def confirm_match(self):
        entry_id = self.wikidata_entries[self.current_index][0]
        any_checked = False
        for i in range(self.stationsList.count()):
            item = self.stationsList.item(i)
            if item.checkState() == Qt.Checked:
                any_checked = True
                station_name = item.text().split(' (')[0]
                matched_station = next((s for s in self.stations_entries if s[2] == station_name), None)
                if matched_station:
                    self.update_wikidata_entry(self.current_index, matched_station[0])
        if any_checked:
            self.cursor.execute("UPDATE wikidata SET \"Linked ITU\" = 1 WHERE \"CPS Wiki ID\" = ?", (entry_id,))
        else:
            self.cursor.execute("UPDATE wikidata SET \"Linked ITU\" = 0 WHERE \"CPS Wiki ID\" = ?", (entry_id,))
        self.conn.commit()
        self.next_entry()

    def no_match(self):
        entry_id = self.wikidata_entries[self.current_index][0]
        self.cursor.execute("UPDATE wikidata SET \"Linked ITU\" = 0 WHERE \"CPS Wiki ID\" = ?", (entry_id,))
        self.conn.commit()
        self.next_entry()

    def next_entry(self):
        self.current_index += 1
        self.show_entry()

    def update_wikidata_entry(self, index, station_id):
        entry_id = self.wikidata_entries[index][0]
        self.cursor.execute("INSERT INTO wikidata_stations_link (\"CPS Wiki ID\", \"CPS Station ID\") VALUES (?, ?)",
                            (entry_id, station_id))
        self.conn.commit()

    def generateMapHTML(self, station_data):
        wikidata_station = station_data[0]
        wikidata_name = wikidata_station[0]
        wikidata_country = wikidata_station[2]
        wikidata_lat = wikidata_station[3]
        wikidata_lon = wikidata_station[4]
        if wikidata_lon is None or wikidata_lat is None:            
            html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>No coordinates</title>
        </head>
        <body>
            <p>No coordinates for wikidata station to show the map</p>
        </body>
        </html>
        """
        else:
            try:
                html_parts = []
                html_parts.append("""
                <!DOCTYPE html>
                <html>
                <head>
                    <title>Full Widget Leaflet Map for Site Link Wizard</title>
                    <meta charset="utf-8" />
                    <link 
                        rel="stylesheet" 
                        href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"
                    />
                    <link
                        rel="stylesheet"
                        href="https://unpkg.com/leaflet.markercluster/dist/MarkerCluster.css"
                    />
                    <link
                        rel="stylesheet"
                        href="https://unpkg.com/leaflet.markercluster/dist/MarkerCluster.Default.css"
                    />
                    <script 
                        src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js">
                    </script>
                    <script
                        src="https://unpkg.com/leaflet.markercluster/dist/leaflet.markercluster.js">
                    </script>
                    <style>
                        body {
                            padding: 0;
                            margin: 0;
                        }
                        html, body, #map {
                            height: 100%;
                            width: 100%;
                        }
                        .custom-cluster-icon {
                            background: radial-gradient(circle, white 25%, transparent 75%);
                            border-radius: 50%;
                            border: 2px solid rgba(0, 0, 0, 0.5);
                            text-align: center;
                            color: black;
                            font-size: 14px;
                            font-weight: bold;
                            width: 40px;
                            height: 40px;
                        }
                        .custom-cluster-icon img {
                            position: absolute;
                            top: 50%;
                            left: 50%;
                            transform: translate(-50%, -50%);
                            width: 30px;
                            height: 30px;
                        }
                        .custom-cluster-icon .cluster-count {
                            position: absolute;
                            top: -10px;
                            right: -10px;
                            background: red;
                            color: white;
                            border-radius: 50%;
                            width: 20px;
                            height: 20px;
                            display: flex;
                            justify-content: center;
                            align-items: center;
                            font-size: 12px;
                        }
                    </style>
                </head>
                <body>
                    <div id="map"></div>
                
                    <script>
                    """)
                html_parts.append(f"""
                        var map = L.map('map', {{attributionControl: false}}).setView([{wikidata_lat}, {wikidata_lon}], 16);
                        """)
                html_parts.append("""
                        var myAttrControl = L.control.attribution().addTo(map);
                        myAttrControl.setPrefix('<a href="https://leafletjs.com/">Leaflet</a>');
                        
                        mapLink = 
                            '<a href="http://openstreetmap.org">OpenStreetMap</a>';
                        L.tileLayer(
                            'http://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png', {
                            attribution: 'Map data by &copy; <a href="http://www.openstreetmap.org/copyright">OpenStreetMap</a>, under <a href="https://opendatacommons.org/licenses/odbl/">ODbL.</a>',
                            maxZoom: 18,
                            }).addTo(map);
                            
                        var redIcon = new L.Icon({
                            iconUrl: 'https://raw.githubusercontent.com/pointhi/leaflet-color-markers/master/img/marker-icon-red.png',
                            shadowUrl: 'https://cdnjs.cloudflare.com/ajax/libs/leaflet/1.3.1/images/marker-shadow.png',
                            iconSize: [25, 41],
                            iconAnchor: [12, 41],
                            popupAnchor: [1, -34],
                            shadowSize: [41, 41]
                        });
                """)

                for name, adm, ctr, lat, lon in station_data[1:]:
                    adm_escaped = adm.replace("'", "&#39;").replace('"', '&quot;')
                    ctr_escaped = ctr.replace("'", "&#39;").replace('"', '&quot;')
                    name_escaped = name.replace("'", "&#39;").replace('"', '&quot;')
                    html_parts.append(
                        f"        L.marker([{lat}, {lon}]).addTo(map).bindPopup('<b>ITU Station: {name_escaped}</b><br>Region country code:<b>{ctr_escaped}</b><br>Responsible administration: <b>{adm_escaped}</b>');")

                html_parts.append(
                    f"        L.marker([{wikidata_lat}, {wikidata_lon}], {{icon: redIcon}}).addTo(map).bindPopup('<b>Wikidata station: {wikidata_name.replace("'", "&#39;").replace('"', '&quot;')}</b><br>Country:<b>{wikidata_country.replace("'", "&#39;").replace('"', '&quot;')}</b>').openPopup();")

                html_parts.append("""
                    </script>
                </body>
                </html>
                """)

                html = ''.join(html_parts)

            except Exception as e:
                html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Error</title>
            </head>
            <body>
                <p>Error generating map: {str(e)}</p>
            </body>
            </html>
            """
        return html
    def showEvent(self, event):
        super().showEvent(event)
        if self.map_html:
            self.mapLayout.removeWidget(self.browser)
            self.browser.deleteLater()
            self.browser = QWebEngineView()
            self.mapLayout.addWidget(self.browser)
            self.browser.setHtml(self.map_html)

    def closeEvent(self, event):
        if self.parent:
            self.parent.animateClosing(self)
            self.parent.showNormal()
            self.parent.animateOpening(
                self.parent, self.parent.desired_width, self.parent.desired_height)
            self.parent.setFocus()
            self.parent.raise_()
            self.parent.activateWindow()
            self.parent.setEnabled(True)


if __name__ == '__main__':
    qt_app = QApplication(sys.argv)
    qt_gui = MainApp()
    sys.exit(qt_app.exec_())
